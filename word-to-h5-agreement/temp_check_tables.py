# -*- coding: utf-8 -*-
from docx import Document

docx_path = r"D:\Seafile\国宝奇妙游小游戏小程序\08 小红书+微信小程序方案\09 条款文件\唐小妹协议条款 final\个人信息第三方共享清单0116.docx"

doc = Document(docx_path)

print(f'段落数: {len(doc.paragraphs)}')
print(f'表格数: {len(doc.tables)}')

print('\n表格内容:')
for i, table in enumerate(doc.tables):
    print(f'\n表格 {i+1}: {len(table.rows)}行 x {len(table.columns)}列')
    for row_idx, row in enumerate(table.rows):
        row_text = [cell.text.strip() for cell in row.cells]
        print(f'  行{row_idx}: {" | ".join(row_text)}')
