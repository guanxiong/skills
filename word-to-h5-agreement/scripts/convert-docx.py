#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档转HTML脚本
严格提取Word内容，生成美观的H5页面
智能识别标题和生效日期
"""

import sys
import os
import re
from pathlib import Path
from datetime import datetime

# 设置UTF-8输出编码（Windows兼容）
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

try:
    from docx import Document
except ImportError:
    print("[错误] 未安装 python-docx")
    print("请运行: pip install python-docx")
    sys.exit(1)

def extract_title_and_date(doc, docx_path):
    """
    智能识别文档标题和生效日期
    """
    title = None
    date = None

    for para in doc.paragraphs:
        text = para.text.strip()

        # 识别标题（第一个非空段落）
        if not title and text:
            title = text

        # 识别生效日期
        if not date and "生效日期" in text:
            # 提取日期部分
            match = re.search(r'(\d{4}年\d{1,2}月\d{1,2}日|\d{4}-\d{1,2}-\d{1,2})', text)
            if match:
                date = match.group(1)
            else:
                # 如果提取不到日期，保留整段
                date = text.replace("生效日期：", "").replace("生效日期:", "").strip()

    # 如果没有提取到标题，使用文件名
    if not title:
        return Path(docx_path).stem, date

    return title, date

def extract_formatted_text(paragraph):
    """
    提取段落文本并保留加粗等格式，包括超链接
    """
    html_parts = []

    # 首先处理普通的run
    for run in paragraph.runs:
        if not run.text:
            continue
        text = run.text
        # 检查是否有加粗
        if run.bold:
            html_parts.append(f"<strong>{text}</strong>")
        # 检查是否有斜体
        elif run.italic:
            html_parts.append(f"<em>{text}</em>")
        # 检查是否有下划线
        elif run.underline:
            html_parts.append(f"<u>{text}</u>")
        else:
            html_parts.append(text)

    # 处理hyperlink（超链接）中的文本
    # python-docx的runs不会包含hyperlink，需要从XML中提取
    for child in paragraph._element.iterchildren():
        # 检查是否是hyperlink元素
        if child.tag.endswith('}hyperlink'):
            # 获取hyperlink中的文本
            hyperlink_text = ''
            for text_elem in child.iter():
                if text_elem.tag.endswith('}t'):
                    if text_elem.text:
                        hyperlink_text += text_elem.text

            if not hyperlink_text:
                continue

            # 检查hyperlink中是否有格式
            is_bold = False
            is_italic = False
            is_underline = False

            # 检查rPr（run properties）中的格式
            for rpr in child.iterchildren():
                if rpr.tag.endswith('}rPr'):
                    for prop in rpr.iterchildren():
                        if prop.tag.endswith('}b') or prop.tag.endswith('}bCs'):
                            is_bold = True
                        elif prop.tag.endswith('}i') or prop.tag.endswith('}iCs'):
                            is_italic = True
                        elif prop.tag.endswith('}u'):
                            is_underline = True

            # 应用格式
            formatted_text = hyperlink_text
            if is_bold:
                formatted_text = f"<strong>{formatted_text}</strong>"
            if is_italic:
                formatted_text = f"<em>{formatted_text}</em>"
            if is_underline:
                formatted_text = f"<u>{formatted_text}</u>"

            html_parts.append(formatted_text)

    # 如果没有提取到内容，回退到paragraph.text
    if not html_parts:
        return paragraph.text

    return ''.join(html_parts)

def convert_table_to_html(table):
    """将Word表格转换为HTML表格"""
    html = '<div class="table-wrapper">\n  <table>\n'

    # 判断是否包含表头（第一行作为表头）
    if len(table.rows) > 0:
        # 表头
        header_row = table.rows[0]
        html += '    <thead>\n      <tr>\n'
        for cell in header_row.cells:
            cell_text = cell.text.strip()
            # 检查是否有加粗格式
            is_bold = False
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.bold:
                        is_bold = True
                        break
            cell_html = f'<strong>{cell_text}</strong>' if is_bold else cell_text
            html += f'        <th>{cell_html}</th>\n'
        html += '      </tr>\n    </thead>\n'

        # 表体
        html += '    <tbody>\n'
        for row in table.rows[1:]:
            html += '      <tr>\n'
            for cell in row.cells:
                cell_text = cell.text.strip()
                html += f'        <td>{cell_text}</td>\n'
            html += '      </tr>\n'
        html += '    </tbody>\n'

    html += '  </table>\n</div>'
    return html

def extract_text_from_docx(docx_path):
    """严格提取Word文档的所有内容，不修改任何文字，并保留加粗格式"""
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph

    doc = Document(docx_path)
    content = []

    # 先识别标题和日期
    title, date = extract_title_and_date(doc, docx_path)

    # 提取所有段落和表格（按照文档顺序）
    title_found = False
    date_found = False

    # 遍历文档body的所有元素
    for element in doc.element.body:
        if isinstance(element, CT_P):
            # 处理段落
            para = Paragraph(element, doc)
            plain_text = para.text

            if not plain_text.strip():
                continue

            # 跳过标题段落
            if not title_found and plain_text.strip() == title:
                title_found = True
                continue

            # 跳过生效日期段落
            if not date_found and ("生效日期" in plain_text or plain_text.strip() == date):
                date_found = True
                continue

            # 提取带格式的文本
            formatted_text = extract_formatted_text(para)

            # 识别标题级别
            content.append({
                'type': 'paragraph',
                'text': formatted_text,
                'plain_text': plain_text,
                'style': para.style.name,
                'level': get_heading_level(para, plain_text)
            })

        elif isinstance(element, CT_Tbl):
            # 处理表格
            table = Table(element, doc)
            # 将表格转换为HTML
            table_html = convert_table_to_html(table)
            content.append({
                'type': 'table',
                'html': table_html
            })

    return content, title, date

def get_heading_level(para, text):
    """智能获取标题级别"""
    style_name = para.style.name

    # 优先使用样式名称
    if 'Heading 1' in style_name:
        return 1
    elif 'Heading 2' in style_name:
        return 2
    elif 'Heading 3' in style_name:
        return 3
    elif 'Heading 4' in style_name:
        return 4
    elif 'Heading 5' in style_name:
        return 5
    elif 'Heading 6' in style_name:
        return 6

    # 根据文本内容智能识别
    text = text.strip()

    # 识别主标题（序号格式：一、二、三、）
    if re.match(r'^[一二三四五六七八九十]+、', text):
        return 2

    # 识别主标题（序号格式：1.、2.、3.、4.、等，注意：不包含 "x.x" 格式）
    if re.match(r'^\d+\.\s+[^\d]', text):
        return 2

    # 处理 List Paragraph 样式（通常是一级标题）
    # 但是：如果以【】开头且长度较短，可能是段落而非标题
    if 'List Paragraph' in style_name:
        if re.match(r'^【', text) or len(text) >= 50:
            return 0  # 普通段落
        return 2

    # 识别子标题（序号格式：2.1、3.1、8.1、等）
    # 注意：需要区分子标题和段落编号（如 "9.1 ...长文本..."）
    # 规则：如果文本长度 <= 30，认为是子标题（h3）；否则是段落（p）
    match = re.match(r'^(\d+\.\d+)\s+', text)
    if match and len(text) <= 30:
        return 3

    # 识别列表标题（序号格式：1、2、3、）
    if re.match(r'^\d+、', text):
        return 3

    # 特殊标题：权限申请使用
    if text == '权限申请使用':
        return 3

    # 识别独立的小标题（需要更严格的条件，避免误判列表项为标题）
    # 排除列表项（如"（1）..."、"1、..."、"·..."、"—..."、"【"等）
    # 优先检查是否包含特殊字符开头
    if (len(text) < 30 and
        text.count('。') == 0 and
        text.count('！') == 0 and
        text.count('？') == 0 and
        text.count('；') == 0 and
        text.count('，') == 0 and
        not re.match(r'^[（\(【\[]', text) and  # 不以（、【或[开头
        not re.match(r'^\d+、', text) and
        not re.match(r'^\d+\.', text) and
        not re.match(r'^·', text) and
        not text.startswith('—') and
        not text.startswith('-')):
        # 只有在不以任何列表符号开头时，才认为是 h2
        return 2

    return 0  # 普通段落

def number_to_chinese(num):
    """将数字转换为中文序号（支持1-100+）"""
    if num <= 10:
        mapping = {
            1: '一', 2: '二', 3: '三', 4: '四', 5: '五',
            6: '六', 7: '七', 8: '八', 9: '九', 10: '十'
        }
        return mapping[num]
    elif num < 20:
        return '十' + number_to_chinese(num - 10) if num > 10 else '十'
    elif num < 100:
        tens = num // 10
        units = num % 10
        result = number_to_chinese(tens) + '十'
        if units > 0:
            result += number_to_chinese(units)
        return result
    else:
        return str(num)

def fix_heading_numbers(content):
    """
    智能修复标题序号，确保连续性和一致性

    处理以下情况：
    1. 第一个标题缺少序号（如"个人信息的收集和使用"，但后面是"2. 信息的存储"）
    2. 序号不连续（如1, 3, 4）
    3. 序号格式不一致（部分有中文序号"一、"，部分有数字序号"1."）
    4. 根据文档主要使用的序号格式进行统一和补齐
    """
    # 只处理段落类型的content，跳过表格类型
    h2_items = [item for item in content if item.get('type') == 'paragraph' and item.get('level') == 2]

    if len(h2_items) < 2:
        return content  # 标题太少，无需修复

    # 统计序号格式（使用plain_text进行匹配）
    has_chinese_num = 0
    has_digit_num = 0
    has_no_num = 0

    for item in h2_items:
        text = item['plain_text'].strip()  # 使用纯文本进行匹配
        if re.match(r'^[一二三四五六七八九十]+、', text):
            has_chinese_num += 1
        elif re.match(r'^\d+\.\s', text):
            has_digit_num += 1
        else:
            has_no_num += 1

    # 决定使用哪种序号格式
    # 如果中文序号更多，使用中文序号；否则使用数字序号
    use_chinese = has_chinese_num > has_digit_num and has_chinese_num > 0
    use_digit = has_digit_num >= has_chinese_num and has_digit_num > 0

    # 如果两种序号都有，按数量多的一方为准
    # 如果都没有序号，默认使用数字序号
    if not use_chinese and not use_digit:
        use_digit = True

    # 根据选择的格式统一和补齐序号
    for i, item in enumerate(h2_items):
        plain_text = item['plain_text'].strip()  # 使用纯文本提取序号
        formatted_text = item['text']  # 使用带格式的文本
        num = i + 1

        # 提取标题文本（去掉已有序号）
        title_text = None

        # 尝试匹配中文序号
        match_cn = re.match(r'^[一二三四五六七八九十]+、\s*(.+)$', plain_text)
        if match_cn:
            title_text = match_cn.group(1)
        else:
            # 尝试匹配数字序号
            match_digit = re.match(r'^\d+\.\s+(.+)$', plain_text)
            if match_digit:
                title_text = match_digit.group(1)
            else:
                # 没有序号，直接使用带格式的全文
                title_text = formatted_text

        # 根据选择的格式设置序号（保留带格式的标题文本）
        if use_chinese:
            cn_num = number_to_chinese(num)
            item['text'] = f"{cn_num}、{title_text}"
        elif use_digit:
            item['text'] = f"{num}. {title_text}"

    return content

def generate_english_filename(title, date):
    """
    根据标题和生效日期生成英文文件名

    Args:
        title: 文档标题（中文）
        date: 生效日期（如 "2026年1月1日" 或 "2026-01-01"）

    Returns:
        英文文件名（如 "privacy-policy-20260101.html"）
    """
    # 根据标题判断协议类型
    title_lower = title.lower()

    if '隐私' in title and '儿童' in title:
        english_name = 'children-privacy'
    elif '隐私' in title:
        english_name = 'privacy-policy'
    elif '用户' in title and ('协议' in title or '服务' in title):
        english_name = 'user-agreement'
    else:
        # 默认使用 document
        english_name = 'document'

    # 格式化日期
    date_str = format_date_for_filename(date)

    return f"{english_name}-{date_str}.html"

def format_date_for_filename(date):
    """
    将日期格式化为文件名友好的格式（YYYYMMDD）

    Args:
        date: 日期字符串（如 "2026年1月1日" 或 "2026-01-01"）

    Returns:
        格式化的日期字符串（如 "20260101"）
    """
    if not date:
        # 默认使用当前月份的1日
        now = datetime.now()
        return f"{now.year}{now.month:02d}01"

    # 匹配 "2026年1月1日" 或 "2026年01月01日"
    match = re.match(r'(\d{4})年(\d{1,2})月(\d{1,2})日', date)
    if match:
        year, month, day = match.groups()
        return f"{year}{month.zfill(2)}{day.zfill(2)}"

    # 匹配 "2026-01-01"
    match = re.match(r'(\d{4})-(\d{1,2})-(\d{1,2})', date)
    if match:
        year, month, day = match.groups()
        return f"{year}{month.zfill(2)}{day.zfill(2)}"

    # 无法解析，返回默认日期
    return '19700101'

def get_document_type_from_filename(filename):
    """
    根据HTML文件名识别文档类型

    Args:
        filename: 文件名（不带扩展名）

    Returns:
        文档类型字符串（user-agreement, privacy-policy, children-privacy等）
    """
    filename_lower = filename.lower()

    # 优先匹配明确的关键词
    if 'user-agreement' in filename_lower or '用户协议' in filename_lower:
        return 'user-agreement'
    elif 'privacy-policy' in filename_lower or '隐私协议' in filename_lower:
        return 'privacy-policy'
    elif 'children-privacy' in filename_lower or '儿童隐私' in filename_lower:
        return 'children-privacy'
    elif 'document' in filename_lower:
        return 'document'
    else:
        # 默认使用user-agreement
        return 'user-agreement'

def get_footer_links(docx_path):
    """
    扫描目录，根据已存在的HTML文件动态生成footer链接

    Args:
        docx_path: Word文件路径

    Returns:
        footer链接HTML字符串
    """
    dir_path = Path(docx_path).parent

    # 定义文档类型到显示名称的映射
    doc_type_names = {
        'user-agreement': '用户协议',
        'privacy-policy': '隐私协议',
        'children-privacy': '儿童隐私保护',
        'document': '第三方清单',
    }

    # 扫描目录中的所有HTML文件
    html_files = list(dir_path.glob("*.html"))

    # 按文档类型分组，只保留最新的版本（如果有多个日期版本）
    latest_files = {}
    for html_file in html_files:
        filename = html_file.stem
        doc_type = get_document_type_from_filename(filename)

        # 如果该类型还没有文件，或者当前文件更新（通过修改时间判断）
        if doc_type not in latest_files:
            latest_files[doc_type] = html_file
        else:
            # 比较修改时间，保留最新的
            if html_file.stat().st_mtime > latest_files[doc_type].stat().st_mtime:
                latest_files[doc_type] = html_file

    # 生成footer链接
    links = []
    for doc_type, html_file in latest_files.items():
        label = doc_type_names.get(doc_type, html_file.stem)
        links.append(f'          <a href="./{html_file.name}">{label}</a>\n')

    # 按照固定顺序排序：用户协议、隐私协议、儿童隐私、其他
    order = ['user-agreement', 'privacy-policy', 'children-privacy', 'document']
    sorted_links = []
    for doc_type in order:
        if doc_type in latest_files:
            html_file = latest_files[doc_type]
            label = doc_type_names[doc_type]
            sorted_links.append(f'          <a href="./{html_file.name}">{label}</a>\n')

    # 添加其他类型
    for doc_type, html_file in latest_files.items():
        if doc_type not in order:
            label = doc_type_names.get(doc_type, html_file.stem)
            sorted_links.append(f'          <a href="./{html_file.name}">{label}</a>\n')

    return ''.join(sorted_links)

def generate_html(content, title, date, docx_path):
    """生成HTML文件，内容严格从Word提取"""
    docx_file = Path(docx_path)
    # 使用英文+日期格式生成文件名
    html_filename = generate_english_filename(title, date)
    html_path = docx_file.parent / html_filename

    # 获取动态footer链接
    footer_links = get_footer_links(docx_path)

    # 智能修复标题序号
    content = fix_heading_numbers(content)

    # 处理日期
    if date:
        date_display = date
    else:
        # 默认使用当前月份的1日
        now = datetime.now()
        date_display = f"{now.year}年{now.month}月1日"

    # CSS 样式（内联）
    css = """<style>
:root {
  --primary-color: #E88A7A;
  --primary-light: #F0A898;
  --primary-dark: #D86A5A;
  --bg-color: #FFF9F7;
  --card-bg: #FFFFFF;
  --text-primary: #3D3A38;
  --text-secondary: #6B6662;
  --text-light: #9A958F;
  --border-color: #F0EBE7;
  --spacing-xs: 8px;
  --spacing-sm: 12px;
  --spacing-md: 16px;
  --spacing-lg: 24px;
  --spacing-xl: 32px;
  --border-radius: 12px;
  --shadow: 0 2px 12px rgba(232, 138, 122, 0.1);
}
* {
  margin: 0;
  padding: 0;
  box-sizing: border-box;
}
html {
  font-size: 16px;
  -webkit-text-size-adjust: 100%;
  -webkit-tap-highlight-color: transparent;
}
body {
  font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'PingFang SC', 'Hiragino Sans GB', 'Microsoft YaHei', 'Helvetica Neue', Helvetica, Arial, sans-serif;
  background: var(--bg-color);
  color: var(--text-primary);
  line-height: 1.8;
  min-height: 100vh;
}
body.in-miniprogram {
  padding-bottom: env(safe-area-inset-bottom);
}
.page-wrapper {
  min-height: 100vh;
  display: flex;
  flex-direction: column;
}
.header {
  background: var(--card-bg);
  padding: var(--spacing-lg) var(--spacing-md);
  border-bottom: 1px solid var(--border-color);
}
.header-content {
  max-width: 800px;
  margin: 0 auto;
}
.header-title {
  color: var(--text-primary);
  font-size: 1.5rem;
  font-weight: 600;
  margin-bottom: var(--spacing-xs);
}
.header-subtitle {
  color: var(--text-light);
  font-size: 0.85rem;
}
.main-content {
  flex: 1;
  padding: var(--spacing-lg) var(--spacing-md);
  max-width: 800px;
  margin: 0 auto;
  width: 100%;
}
.content-card {
  background: var(--card-bg);
  border-radius: var(--border-radius);
  box-shadow: var(--shadow);
  padding: var(--spacing-lg) var(--spacing-md);
  overflow: hidden;
}
.content-card h1 {
  font-size: 1.4rem;
  font-weight: 700;
  color: var(--text-primary);
  margin-bottom: var(--spacing-md);
  padding-bottom: var(--spacing-sm);
  border-bottom: 2px dashed var(--border-color);
}
.content-card h2 {
  font-size: 1.15rem;
  font-weight: 700;
  color: var(--text-primary);
  margin-top: var(--spacing-lg);
  margin-bottom: var(--spacing-sm);
  padding-left: 12px;
  border-left: 4px solid var(--primary-color);
}
.content-card h3 {
  font-size: 1.05rem;
  font-weight: 700;
  color: var(--text-primary);
  margin-top: var(--spacing-md);
  margin-bottom: var(--spacing-xs);
}
.content-card p {
  margin-bottom: var(--spacing-sm);
  color: var(--text-secondary);
  text-align: justify;
  line-height: 1.9;
}
.content-card strong,
.content-card b {
  color: var(--primary-dark);
  font-weight: 600;
}
.content-card a {
  color: var(--primary-color);
  text-decoration: underline;
  text-underline-offset: 2px;
}
.content-card ol,
.content-card ul {
  margin: var(--spacing-sm) 0;
  padding-left: var(--spacing-lg);
  color: var(--text-secondary);
  line-height: 1.9;
}
.content-card ol li,
.content-card ul li {
  margin-bottom: var(--spacing-xs);
}
.content-card ol {
  list-style: decimal;
}
.content-card ol li::marker {
  color: var(--primary-color);
  font-weight: 600;
}
.content-card ul {
  list-style: disc;
}
.content-card ul li::marker {
  color: var(--primary-light);
}
.emphasis-box {
  background: #FAF3F1;
  border-left: 3px solid var(--primary-color);
  border-radius: 4px;
  padding: var(--spacing-md);
  margin: var(--spacing-md) 0;
}
.emphasis-box p {
  margin-bottom: 0;
  color: var(--text-primary);
}
.table-wrapper {
  margin: var(--spacing-lg) 0;
  overflow-x: auto;
  -webkit-overflow-scrolling: touch;
}
.content-card table {
  width: 100%;
  border-collapse: collapse;
  background: #fff;
  border-radius: 8px;
  overflow: hidden;
  box-shadow: 0 2px 8px rgba(0, 0, 0, 0.06);
  min-width: 600px;
}
.content-card thead {
  background: var(--primary-color);
}
.content-card thead th {
  color: #fff;
  font-weight: 600;
  padding: var(--spacing-md);
  text-align: left;
  border-bottom: 2px solid var(--primary-dark);
}
.content-card tbody tr {
  border-bottom: 1px solid var(--border-color);
  transition: background-color 0.2s;
}
.content-card tbody tr:last-child {
  border-bottom: none;
}
.content-card tbody tr:hover {
  background: #FAF9F8;
}
.content-card td {
  padding: var(--spacing-md);
  color: var(--text-secondary);
  line-height: 1.6;
}
.content-card td a {
  color: var(--primary-color);
  text-decoration: none;
  word-break: break-all;
}
.content-card td a:hover {
  text-decoration: underline;
}
.footer {
  background: var(--card-bg);
  border-top: 1px solid var(--border-color);
  padding: var(--spacing-lg) var(--spacing-md);
  margin-top: var(--spacing-lg);
}
.footer-content {
  max-width: 800px;
  margin: 0 auto;
  text-align: center;
}
.footer-links {
  display: flex;
  flex-wrap: wrap;
  justify-content: center;
  gap: var(--spacing-md);
  margin-bottom: var(--spacing-md);
}
.footer-links a {
  color: var(--text-secondary);
  text-decoration: none;
  font-size: 0.85rem;
  transition: color 0.2s;
}
.footer-links a:hover {
  color: var(--primary-color);
}
.footer-copyright {
  color: var(--text-light);
  font-size: 0.8rem;
}
.back-top {
  position: fixed;
  right: var(--spacing-md);
  bottom: var(--spacing-md);
  width: 44px;
  height: 44px;
  background: var(--primary-color);
  color: #fff;
  border: none;
  border-radius: 50%;
  box-shadow: 0 4px 12px rgba(232, 138, 122, 0.25);
  cursor: pointer;
  opacity: 0;
  visibility: hidden;
  transition: all 0.25s;
  z-index: 99;
  display: flex;
  align-items: center;
  justify-content: center;
  font-size: 1rem;
}
.back-top.visible {
  opacity: 1;
  visibility: visible;
}
.back-top:hover {
  background: var(--primary-dark);
  transform: translateY(-2px);
}
@media (min-width: 768px) {
  html {
    font-size: 16px;
  }
  .header {
    padding: var(--spacing-xl) var(--spacing-lg);
  }
  .header-title {
    font-size: 1.5rem;
  }
  .main-content {
    padding: var(--spacing-xl);
  }
  .content-card {
    padding: var(--spacing-xl);
  }
  .content-card h1 {
    font-size: 1.5rem;
  }
  .content-card h2 {
    font-size: 1.2rem;
  }
  .footer-links a {
    font-size: 0.95rem;
  }
}
@media (min-width: 1200px) {
  .main-content {
    padding: var(--spacing-xl) 0;
  }
  .content-card {
    padding: 40px;
  }
}
@media (max-width: 374px) {
  html {
    font-size: 14px;
  }
  .header {
    padding: var(--spacing-md) var(--spacing-sm);
  }
  .main-content {
    padding: var(--spacing-md) var(--spacing-sm);
  }
  .content-card {
    padding: var(--spacing-md) var(--spacing-sm);
    border-radius: var(--spacing-sm);
  }
}
@media (min-width: 376px) and (max-width: 767px) {
  body.in-miniprogram .content-card {
    border-radius: 0;
  }
}
@media print {
  .header,
  .footer,
  .back-top {
    display: none;
  }
  .content-card {
    box-shadow: none;
    border: none;
  }
  body {
    background: #fff;
  }
}
@keyframes fadeIn {
  from {
    opacity: 0;
    transform: translateY(10px);
  }
  to {
    opacity: 1;
    transform: translateY(0);
  }
}
.main-content {
  animation: fadeIn 0.4s ease-out;
}
::-webkit-scrollbar {
  width: 6px;
  height: 6px;
}
::-webkit-scrollbar-track {
  background: var(--bg-color);
}
::-webkit-scrollbar-thumb {
  background: var(--primary-light);
  border-radius: 3px;
}
::-webkit-scrollbar-thumb:hover {
  background: var(--primary-color);
}
</style>
"""

    # HTML 模板
    html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0, user-scalable=no, viewport-fit=cover">
  <meta name="format-detection" content="telephone=no, email=no">
  <meta name="apple-mobile-web-app-capable" content="yes">
  <meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">
  <title>{title}</title>
  {css}
</head>
<body>
  <div class="page-wrapper">
    <header class="header">
      <div class="header-content">
        <h1 class="header-title">{title}</h1>
        <p class="header-subtitle">生效日期：{date_display}</p>
      </div>
    </header>

    <main class="main-content">
      <article class="content-card">
"""

    # 严格添加Word文档内容，不修改任何文字
    for item in content:
        # 处理表格类型
        if item.get('type') == 'table':
            html_content += f'        {item["html"]}\n'
            continue

        # 处理段落类型
        text = item['text']
        level = item['level']

        if level == 1:
            html_content += f'        <h1>{text}</h1>\n'
        elif level == 2:
            html_content += f'        <h2>{text}</h2>\n'
        elif level == 3:
            html_content += f'        <h3>{text}</h3>\n'
        elif level == 4:
            html_content += f'        <h4>{text}</h4>\n'
        elif level == 5:
            html_content += f'        <h5>{text}</h5>\n'
        elif level == 6:
            html_content += f'        <h6>{text}</h6>\n'
        else:
            # 普通段落，保持原样
            html_content += f'        <p>{text}</p>\n'

    # 结束HTML
    html_content += """      </article>
    </main>

    <footer class="footer">
      <div class="footer-content">
        <nav class="footer-links">
""" + footer_links + """        </nav>
        <p class="footer-copyright">上海东桓文化科技有限公司 © 2026</p>
      </div>
    </footer>

    <button class="back-top" id="backTop" aria-label="返回顶部">▲</button>
  </div>

  <script>
    const backTop = document.getElementById('backTop');
    function toggleBackTop() {
      if (window.scrollY > 300) {
        backTop.classList.add('visible');
      } else {
        backTop.classList.remove('visible');
      }
    }
    backTop.addEventListener('click', () => {
      window.scrollTo({ top: 0, behavior: 'smooth' });
    });
    window.addEventListener('scroll', toggleBackTop);
    toggleBackTop();

    const isMiniprogram = /miniprogram/i.test(navigator.userAgent) ||
                         window.__wxjs_environment === 'miniprogram';
    if (isMiniprogram) document.body.classList.add('in-miniprogram');
  </script>
</body>
</html>
"""

    # 写入HTML文件
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html_content)

    return html_path

def main():
    if len(sys.argv) != 2:
        print("使用方法: python convert-docx.py <word文件路径>")
        sys.exit(1)

    docx_path = sys.argv[1]

    if not os.path.exists(docx_path):
        print(f"[错误] 文件不存在: {docx_path}")
        sys.exit(1)

    if not docx_path.lower().endswith('.docx'):
        print(f"[错误] 文件格式错误，请提供 .docx 文件: {docx_path}")
        sys.exit(1)

    try:
        print(f"[开始] 正在读取Word文档...")
        content, title, date = extract_text_from_docx(docx_path)
        print(f"[成功] 提取了 {len(content)} 个段落")
        print(f"[信息] 标题：{title}")
        print(f"[信息] 生效日期：{date}")

        html_path = generate_html(content, title, date, docx_path)

        print(f"\n{'='*60}")
        print(f"[完成] HTML文件已生成")
        print(f"{'='*60}")
        print(f"Word文件: {docx_path}")
        print(f"HTML文件: {html_path.absolute()}")
        print(f"{'='*60}\n")
        print(f"请在浏览器中访问以下地址查看效果：")
        print(f"file:///{html_path.absolute().as_posix()}")
        print(f"\n提示：Windows用户可以直接双击HTML文件打开")

    except Exception as e:
        print(f"[错误] 转换失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == '__main__':
    main()
